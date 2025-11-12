from flask import Flask, render_template_string, request, send_file, jsonify, session
from flask_session import Session
import os, uuid, io
from datetime import datetime
import pytz
from openai import OpenAI
from docx import Document
import tempfile

# ---------------- CONFIG ----------------
APP_NAME = "FOSCHI IA WEB"
CREADOR = "Gustavo Enrique Foschi"
DATA_DIR = "data"
STATIC_DIR = "static"

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

# ---------------- API KEYS ----------------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY)

# ---------------- APP ----------------
app = Flask(__name__)
app.secret_key = "FoschiWebKey"
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

# ---------------- UTILIDADES ----------------
def fecha_hora_en_es():
    tz = pytz.timezone("America/Argentina/Buenos_Aires")
    ahora = datetime.now(tz)
    return ahora.strftime("%d/%m/%Y %H:%M")

# ---------------- RUTA: PÁGINA PRINCIPAL ----------------
@app.route("/")
def index():
    if "usuario_id" not in session:
        session["usuario_id"] = str(uuid.uuid4())
    return render_template_string(HTML_TEMPLATE, APP_NAME=APP_NAME, usuario_id=session["usuario_id"])

# ---------------- RUTA: SUBIR VIDEO Y TRANSCRIBIR ----------------
@app.route("/subir_video", methods=["POST"])
def subir_video():
    if "video" not in request.files:
        return jsonify({"error": "No se subió ningún archivo"}), 400

    video = request.files["video"]
    nombre_original = os.path.splitext(video.filename)[0]
    if not nombre_original:
        nombre_original = f"video_{uuid.uuid4().hex}"

    # Guardar temporalmente el video
    video_path = os.path.join(tempfile.gettempdir(), f"{nombre_original}.mp4")
    video.save(video_path)

    try:
        # Transcribir con Whisper API (OpenAI)
        with open(video_path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="text"
            )

        texto = transcript.strip() if transcript else "No se pudo obtener transcripción."

        # Crear archivo Word
        doc = Document()
        doc.add_heading(f"Transcripción de {nombre_original}", level=1)
        doc.add_paragraph(texto)
        doc_path = os.path.join(tempfile.gettempdir(), f"{nombre_original}.docx")
        doc.save(doc_path)

        # Preparar descarga
        return_data = io.BytesIO()
        with open(doc_path, "rb") as f:
            return_data.write(f.read())
        return_data.seek(0)

        # Borrar archivos temporales
        try:
            os.remove(video_path)
            os.remove(doc_path)
        except:
            pass

        # Enviar el archivo Word al usuario
        return send_file(
            return_data,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{nombre_original}.docx"
        )

    except Exception as e:
        return jsonify({"error": f"Error en transcripción: {e}"}), 500


# ---------------- HTML MODERNO ----------------
HTML_TEMPLATE = """
<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <title>{{APP_NAME}} — Transcriptor</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body {
      background: linear-gradient(135deg, #000000, #1a1a1a);
      color: #fff;
      font-family: 'Segoe UI', Roboto, sans-serif;
      text-align: center;
      margin: 0;
      padding: 0;
    }
    h1 { margin-top: 20px; }
    .container {
      margin: 40px auto;
      padding: 20px;
      background: #111;
      border-radius: 20px;
      width: 90%;
      max-width: 500px;
      box-shadow: 0 0 20px #00ffff33;
    }
    input[type=file] {
      display: none;
    }
    label {
      display: inline-block;
      padding: 12px 25px;
      background: #00ffff;
      color: #000;
      border-radius: 30px;
      cursor: pointer;
      font-weight: bold;
      transition: 0.3s;
    }
    label:hover { background: #00cccc; }
    #mensaje {
      margin-top: 15px;
      color: #aaa;
    }
    #loading {
      display: none;
      margin-top: 20px;
      font-style: italic;
      color: #0ff;
    }
    button {
      margin-top: 20px;
      padding: 10px 20px;
      background: #00ffff;
      border: none;
      color: #000;
      font-weight: bold;
      border-radius: 25px;
      cursor: pointer;
      transition: 0.3s;
    }
    button:hover { background: #00cccc; }
  </style>
</head>
<body>
  <h1>🎬 FOSCHI IA — Transcriptor de Video</h1>
  <div class="container">
    <form id="formVideo" enctype="multipart/form-data">
      <label for="video">📁 Elegí un video</label>
      <input type="file" id="video" name="video" accept="video/*" required>
      <div id="mensaje">Ningún archivo seleccionado</div>
      <button type="submit">Transcribir y Descargar Word</button>
    </form>
    <div id="loading">⏳ Transcribiendo audio... Esto puede tardar unos minutos</div>
  </div>

  <script>
  const form = document.getElementById('formVideo');
  const input = document.getElementById('video');
  const msg = document.getElementById('mensaje');
  const loading = document.getElementById('loading');

  input.addEventListener('change', () => {
    if (input.files.length > 0) {
      msg.textContent = '🎞️ ' + input.files[0].name;
    } else {
      msg.textContent = 'Ningún archivo seleccionado';
    }
  });

  form.addEventListener('submit', async (e) => {
    e.preventDefault();
    if (input.files.length === 0) {
      alert('Subí un video primero');
      return;
    }

    loading.style.display = 'block';
    msg.textContent = 'Procesando...';

    const formData = new FormData();
    formData.append('video', input.files[0]);

    const response = await fetch('/subir_video', { method: 'POST', body: formData });
    loading.style.display = 'none';

    if (response.ok) {
      const blob = await response.blob();
      const link = document.createElement('a');
      const fileName = input.files[0].name.replace(/\.[^/.]+$/, "") + ".docx";
      link.href = window.URL.createObjectURL(blob);
      link.download = fileName;
      link.click();
      msg.textContent = '✅ Transcripción completada y descargada';
      input.value = '';
    } else {
      const err = await response.json();
      alert('Error: ' + (err.error || 'No se pudo procesar el video.'));
      msg.textContent = '❌ Error en la transcripción';
    }
  });
  </script>
</body>
</html>
"""

# ---------------- MAIN ----------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
